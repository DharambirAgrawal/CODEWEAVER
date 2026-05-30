#!/usr/bin/env python3
# src/execify/probes/probe_docx.py
# Run: python probe_docx.py <path_to_docx>
# Returns JSON to stdout.
import sys
import json
try:
    from docx import Document
except ImportError:
    print(json.dumps({"error": "python-docx not installed"}))
    sys.exit(1)

if len(sys.argv) < 2:
    print(json.dumps({"error": "Usage: probe_docx.py <path>"}))
    sys.exit(1)

path = sys.argv[1]
try:
    doc = Document(path)
except Exception as e:
    print(json.dumps({"error": str(e)}))
    sys.exit(1)

paras = [p for p in doc.paragraphs if p.text.strip()]
words = sum(len(p.text.split()) for p in paras)
headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]

print(json.dumps({
    "paragraph_count": len(paras),
    "word_count": words,
    "heading_count": len(headings),
    "first_paragraph": paras[0].text[:80] if paras else "",
}))
